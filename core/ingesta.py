"""
INGESTA — Extracción de texto crudo de documentos

Estrategia por tipo de archivo:
  PDF nativo      → PyMuPDF (texto) + pdfplumber (tablas)
  PDF escaneado   → pypdfium2/PyMuPDF/pdf2image → OCR + calidad de imagen
  Excel .xlsx/.xls→ python-calamine, fallback openpyxl
  Word .docx      → python-docx
  Office legacy   → LibreOffice headless → PDF → ingesta recursiva
  Imagen          → OpenCV/deskew + OCR

Devuelve ResultadoIngesta con texto, fuente y metadatos.
"""

import io
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from PIL import Image
import openpyxl

from config import OCR_ENGINE, OCR_LANGUAGES, OCR_GPU, OCR_MIN_QUALITY, TESSERACT_LANG, MIN_TEXTO_PDF_COMPLETO, OCR_MAX_PAGINAS
from core.pdf_desbloqueo import preparar_pdf_para_ingesta

# EasyOCR se carga lazy para no bloquear el arranque del servidor
_reader = None
_paddle_reader = None


def _get_reader():
    """Lazy loader del lector EasyOCR."""
    global _reader
    if _reader is None:
        import easyocr
        _reader = easyocr.Reader(OCR_LANGUAGES, gpu=OCR_GPU, verbose=False)
    return _reader


def _ocr_easyocr(imagen: Image.Image) -> str:
    """OCR con EasyOCR."""
    reader = _get_reader()
    import numpy as np
    img_array = np.array(imagen)
    resultados = reader.readtext(img_array, detail=0, paragraph=True)
    return "\n".join(resultados)


def _ocr_tesseract(imagen: Image.Image) -> str:
    """OCR con Tesseract local."""
    import pytesseract
    return pytesseract.image_to_string(imagen, lang=TESSERACT_LANG)


def _get_paddle_reader():
    """Lazy loader de PaddleOCR. Puede descargar modelos en el primer uso."""
    global _paddle_reader
    if _paddle_reader is None:
        from paddleocr import PaddleOCR
        _paddle_reader = PaddleOCR(lang="es", use_doc_orientation_classify=True)
    return _paddle_reader


def _ocr_paddle(imagen: Image.Image) -> str:
    """OCR con PaddleOCR, tolerando diferencias menores entre versiones."""
    reader = _get_paddle_reader()
    import numpy as np
    img_array = np.array(imagen.convert("RGB"))
    resultado = reader.predict(img_array)
    partes: list[str] = []
    for item in resultado:
        if isinstance(item, dict):
            textos = item.get("rec_texts") or item.get("texts") or []
            partes.extend(str(t) for t in textos if str(t).strip())
        elif isinstance(item, list):
            for linea in item:
                if isinstance(linea, (list, tuple)) and len(linea) >= 2:
                    dato = linea[1]
                    if isinstance(dato, (list, tuple)) and dato:
                        partes.append(str(dato[0]))
    return "\n".join(partes)


def _leer_ocr(imagen: Image.Image, advertencias: list[str]) -> tuple[str, str]:
    """
    Ejecuta OCR con proveedor configurable.

    OCR_ENGINE:
      - easyocr: usa EasyOCR
      - tesseract: usa Tesseract
      - paddleocr: usa PaddleOCR
      - auto: intenta EasyOCR y cae a Tesseract si falla
    """
    engine = OCR_ENGINE
    if engine not in {"easyocr", "tesseract", "paddleocr", "auto"}:
        advertencias.append(f"OCR_ENGINE inválido '{OCR_ENGINE}', usando easyocr")
        engine = "easyocr"

    if engine == "tesseract":
        return _ocr_tesseract(imagen), "ocr_tesseract"
    if engine == "paddleocr":
        return _ocr_paddle(imagen), "ocr_paddleocr"

    try:
        return _ocr_easyocr(imagen), "ocr_easyocr"
    except Exception as e:
        if engine != "auto":
            raise
        advertencias.append(f"EasyOCR falló ({e}); usando Tesseract")
        try:
            return _ocr_tesseract(imagen), "ocr_tesseract"
        except Exception as e2:
            advertencias.append(f"Tesseract falló ({e2}); usando PaddleOCR")
            return _ocr_paddle(imagen), "ocr_paddleocr"


# ─── TIPOS ────────────────────────────────────────────────────────────────────

@dataclass
class ResultadoIngesta:
    texto: str
    fuente: str           # 'pdf_nativo' | 'ocr_easyocr' | 'ocr_tesseract' | 'excel' | 'vacio'
    paginas: int
    ocr_usado: bool
    calidad_imagen: float  # 0.0–1.0
    advertencias: list[str] = field(default_factory=list)
    errores: list[str] = field(default_factory=list)


# ─── HELPERS ──────────────────────────────────────────────────────────────────

def _calcular_calidad_imagen(image: Image.Image) -> float:
    """
    Heurística de calidad: resolución + contraste.
    Devuelve 0.0–1.0.
    """
    width, height = image.size
    pixels = width * height

    # Resolución: >= 1 MP es buena
    score_resolucion = min(pixels / 1_000_000, 1.0)

    # Contraste: desviación estándar de luminancia
    try:
        gris = image.convert("L")
        import statistics
        pixeles = list(gris.getdata())
        std = statistics.stdev(pixeles[:10000])  # muestra
        score_contraste = min(std / 80.0, 1.0)
    except Exception:
        score_contraste = 0.5

    return round(score_resolucion * 0.4 + score_contraste * 0.6, 3)


def _preprocesar_imagen_ocr(image: Image.Image, advertencias: list[str]) -> Image.Image:
    """Normaliza imagen para OCR: orientación, contraste, ruido y binarización."""
    try:
        import cv2
        import numpy as np

        img = np.array(image.convert("RGB"))
        gris = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

        alto, ancho = gris.shape[:2]
        lado_mayor = max(alto, ancho)
        if lado_mayor < 1600:
            escala = 1600 / lado_mayor
            gris = cv2.resize(gris, None, fx=escala, fy=escala, interpolation=cv2.INTER_CUBIC)

        try:
            from deskew import determine_skew
            angulo = determine_skew(gris)
            if angulo and abs(angulo) > 0.5:
                h, w = gris.shape[:2]
                centro = (w // 2, h // 2)
                matriz = cv2.getRotationMatrix2D(centro, angulo, 1.0)
                gris = cv2.warpAffine(
                    gris,
                    matriz,
                    (w, h),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE,
                )
                advertencias.append(f"Imagen enderezada ({angulo:.2f} grados)")
        except Exception as e:
            advertencias.append(f"No se pudo enderezar imagen: {e}")

        gris = cv2.fastNlMeansDenoising(gris, None, 10, 7, 21)
        gris = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8)).apply(gris)
        binaria = cv2.adaptiveThreshold(
            gris,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            11,
        )
        return Image.fromarray(binaria)
    except Exception as e:
        advertencias.append(f"Preprocesado OpenCV no disponible: {e}")
        return image


def _leer_bytes(fileobj) -> bytes:
    fileobj.seek(0)
    return fileobj.read()


def _texto_suficiente(texto: str, minimo: int = 100) -> bool:
    return len((texto or "").strip()) >= minimo


def _extraer_texto_pdf_nativo(fileobj) -> tuple[str, int, list[str]]:
    """Extrae texto de PDF nativo con PyMuPDF y tablas con pdfplumber."""
    partes = []
    advertencias = []
    paginas = 0

    pdf_bytes = _leer_bytes(fileobj)

    try:
        import fitz
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            paginas = len(doc)
            for pagina in doc:
                partes.append(pagina.get_text("text") or "")
    except Exception as e:
        advertencias.append(f"PyMuPDF falló ({e}); intentando pdfplumber")

    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            paginas = max(paginas, len(pdf.pages))
            for i, page in enumerate(pdf.pages):
                texto_pagina = page.extract_text() or ""
                if texto_pagina and not _texto_suficiente("\n".join(partes), 60):
                    partes.append(texto_pagina)

                # Extraer tablas si hay
                try:
                    tablas = page.extract_tables()
                    for tabla in tablas:
                        for fila in tabla:
                            if fila:
                                partes.append("\t".join(c or "" for c in fila))
                except Exception as e:
                    advertencias.append(f"Página {i+1}: error extrayendo tablas: {e}")
    except Exception as e:
        advertencias.append(f"pdfplumber falló ({e})")

    return "\n".join(partes), paginas, advertencias


def _renderizar_pdf_a_imagenes(pdf_bytes: bytes, advertencias: list[str]) -> list[Image.Image]:
    """Renderiza PDF a imágenes usando pypdfium2, PyMuPDF o pdf2image."""
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(pdf_bytes)
        imagenes = []
        for i in range(len(pdf)):
            page = pdf[i]
            pil = page.render(scale=200 / 72).to_pil()
            imagenes.append(pil.convert("RGB"))
        return imagenes
    except Exception as e:
        advertencias.append(f"pypdfium2 falló ({e}); intentando PyMuPDF")

    try:
        import fitz
        imagenes = []
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(200 / 72, 200 / 72), alpha=False)
                imagen = Image.open(io.BytesIO(pix.tobytes("png")))
                imagenes.append(imagen.convert("RGB"))
        return imagenes
    except Exception as e:
        advertencias.append(f"PyMuPDF render falló ({e}); intentando pdf2image")

    from pdf2image import convert_from_bytes
    return [img.convert("RGB") for img in convert_from_bytes(pdf_bytes, dpi=200)]


def _extraer_texto_ocr(fileobj, nombre_archivo: str, max_paginas: int | None = None) -> tuple[str, int, float, list[str], str]:
    """Extrae texto de PDF escaneado o imagen usando preprocesado y OCR."""
    advertencias = []
    textos = []
    calidades = []
    limite = max_paginas if max_paginas is not None else OCR_MAX_PAGINAS

    nombre_lower = nombre_archivo.lower()
    try:
        if nombre_lower.endswith(".pdf"):
            imagenes = _renderizar_pdf_a_imagenes(_leer_bytes(fileobj), advertencias)
            paginas = len(imagenes)
            if limite > 0 and paginas > limite:
                advertencias.append(f"OCR limitado a {limite} de {paginas} páginas")
                imagenes = imagenes[:limite]
        else:
            fileobj.seek(0)
            imagen = Image.open(fileobj)
            imagenes = [imagen.convert("RGB")]
            paginas = 1
    except Exception as e:
        return "", 0, 0.0, [f"Error preparando imagen para OCR: {e}"], "vacio"

    fuente_ocr = "ocr_easyocr"

    for i, imagen in enumerate(imagenes):
        calidad = _calcular_calidad_imagen(imagen)
        calidades.append(calidad)

        if calidad < OCR_MIN_QUALITY:
            advertencias.append(f"Página {i+1}: calidad de imagen baja ({calidad:.2f})")

        imagen_ocr = _preprocesar_imagen_ocr(imagen, advertencias)
        texto, fuente_ocr = _leer_ocr(imagen_ocr, advertencias)
        textos.append(texto)

    calidad_media = round(sum(calidades) / len(calidades), 3) if calidades else 0.0
    return "\n".join(textos), paginas, calidad_media, advertencias, fuente_ocr


def _es_pdf_nativo(fileobj) -> bool:
    """Detecta si un PDF tiene texto vectorial o es escaneado."""
    pdf_bytes = _leer_bytes(fileobj)
    try:
        import fitz
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            if len(doc) == 0:
                return False
            texto = doc[0].get_text("text") or ""
            return _texto_suficiente(texto)
    except Exception:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            if not pdf.pages:
                return False
            primera = pdf.pages[0]
            texto = primera.extract_text() or ""
            return _texto_suficiente(texto)


def _extraer_texto_excel(fileobj) -> tuple[str, int, list[str]]:
    """Extrae texto de Excel moderno y legacy con calamine; fallback openpyxl."""
    advertencias: list[str] = []
    partes: list[str] = []

    try:
        from python_calamine import load_workbook
        fileobj.seek(0)
        wb = load_workbook(fileobj)
        for sheet_name in wb.sheet_names:
            sheet = wb.get_sheet_by_name(sheet_name)
            for row in sheet.iter_rows():
                fila = "\t".join("" if c is None else str(c) for c in row)
                if fila.strip():
                    partes.append(fila)
        return "\n".join(partes), len(wb.sheet_names), advertencias
    except Exception as e:
        advertencias.append(f"python-calamine falló ({e}); intentando openpyxl")

    fileobj.seek(0)
    wb = openpyxl.load_workbook(fileobj, data_only=True)
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            fila = "\t".join(str(c) if c is not None else "" for c in row)
            if fila.strip():
                partes.append(fila)
    return "\n".join(partes), len(wb.worksheets), advertencias


def _buscar_soffice() -> str | None:
    """Localiza LibreOffice en PATH o rutas típicas de Windows."""
    candidates = [
        os.getenv("SOFFICE_PATH"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def _convertir_legacy_con_libreoffice(fileobj, nombre_archivo: str) -> tuple[bytes | None, list[str], list[str]]:
    """Convierte formatos legacy Office a PDF usando LibreOffice headless."""
    advertencias: list[str] = []
    errores: list[str] = []
    soffice = _buscar_soffice()
    if not soffice:
        return None, advertencias, ["LibreOffice no está disponible para convertir documentos legacy."]

    sufijo = Path(nombre_archivo).suffix or ".bin"
    with tempfile.TemporaryDirectory() as tmp:
        entrada = Path(tmp) / f"entrada{sufijo}"
        salida = Path(tmp) / "entrada.pdf"
        entrada.write_bytes(_leer_bytes(fileobj))
        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            tmp,
            str(entrada),
        ]
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        except Exception as e:
            return None, advertencias, [f"Error ejecutando LibreOffice: {e}"]
        if proc.returncode != 0 or not salida.exists():
            detalle = (proc.stderr or proc.stdout or "").strip()
            return None, advertencias, [f"LibreOffice no pudo convertir {nombre_archivo}: {detalle}"]
        advertencias.append(f"Documento legacy convertido con LibreOffice: {nombre_archivo}")
        return salida.read_bytes(), advertencias, errores


# ─── FUNCIÓN PRINCIPAL ────────────────────────────────────────────────────────

def ingestar(
    fileobj,
    nombre_archivo: str,
    *,
    trabajadores: list[dict] | None = None,
    trabajador_id: str | None = None,
    pdf_password: str | None = None,
    claves_perfil: dict | None = None,
) -> ResultadoIngesta:
    """
    Extrae texto de un documento laboral.

    :param fileobj:        Objeto file-like (bytes en memoria)
    :param nombre_archivo: Nombre del archivo (para detectar extensión)
    :param trabajadores:   Catálogo para desbloqueo PDF e identidad
    :param trabajador_id:  Trabajador prioritario para claves PDF
    :param pdf_password:   Contraseña manual opcional
    :param claves_perfil:  {nif, nss} del perfil de sesión
    :returns:              ResultadoIngesta con texto y metadatos
    """
    nombre_lower = nombre_archivo.lower()
    advertencias: list[str] = []
    errores: list[str] = []

    # ── Excel ──
    if nombre_lower.endswith((".xlsx", ".xls")):
        try:
            texto, paginas, adv = _extraer_texto_excel(fileobj)
            return ResultadoIngesta(
                texto=texto,
                fuente="excel",
                paginas=paginas,
                ocr_usado=False,
                calidad_imagen=1.0,
                advertencias=advertencias + adv,
            )
        except Exception as e:
            advertencias.append(f"Lectura Excel falló ({e}); intentando conversión LibreOffice")
            pdf_bytes, adv, err = _convertir_legacy_con_libreoffice(fileobj, nombre_archivo)
            if pdf_bytes:
                convertido = ingestar(io.BytesIO(pdf_bytes), f"{nombre_archivo}.pdf")
                convertido.advertencias[:0] = advertencias + adv
                return convertido
            return ResultadoIngesta(
                texto="", fuente="vacio", paginas=0, ocr_usado=False,
                calidad_imagen=0.0, advertencias=advertencias + adv, errores=[f"Error leyendo Excel: {e}", *err],
            )

    # ── Word moderno ──
    if nombre_lower.endswith(".docx"):
        try:
            from docx import Document

            fileobj.seek(0)
            doc = Document(fileobj)
            partes = [p.text for p in doc.paragraphs if p.text.strip()]
            for tabla in doc.tables:
                for fila in tabla.rows:
                    celdas = [celda.text.strip() for celda in fila.cells]
                    if any(celdas):
                        partes.append("\t".join(celdas))

            return ResultadoIngesta(
                texto="\n".join(partes),
                fuente="docx",
                paginas=1,
                ocr_usado=False,
                calidad_imagen=1.0,
                advertencias=advertencias,
            )
        except Exception as e:
            return ResultadoIngesta(
                texto="", fuente="vacio", paginas=0, ocr_usado=False,
                calidad_imagen=0.0, errores=[f"Error leyendo DOCX: {e}"],
            )

    # ── Office legacy / OpenDocument ──
    if nombre_lower.endswith((".doc", ".rtf", ".odt")):
        pdf_bytes, adv, err = _convertir_legacy_con_libreoffice(fileobj, nombre_archivo)
        if pdf_bytes:
            convertido = ingestar(io.BytesIO(pdf_bytes), f"{nombre_archivo}.pdf")
            convertido.advertencias[:0] = advertencias + adv
            return convertido
        return ResultadoIngesta(
            texto="", fuente="vacio", paginas=0, ocr_usado=False,
            calidad_imagen=0.0, advertencias=advertencias + adv, errores=err,
        )

    # ── PDF ──
    if nombre_lower.endswith(".pdf"):
        pdf_bytes = _leer_bytes(fileobj)
        pdf_bytes, adv_pdf, err_pdf = preparar_pdf_para_ingesta(
            pdf_bytes,
            trabajadores=trabajadores,
            trabajador_id=trabajador_id,
            pdf_password=pdf_password,
            claves_perfil=claves_perfil,
        )
        advertencias.extend(adv_pdf)
        if err_pdf:
            return ResultadoIngesta(
                texto="",
                fuente="vacio",
                paginas=0,
                ocr_usado=False,
                calidad_imagen=0.0,
                advertencias=advertencias,
                errores=err_pdf,
            )

        buffer = io.BytesIO(pdf_bytes)

        # Siempre intentar texto nativo; si es insuficiente, complementar con OCR.
        texto_nativo = ""
        paginas = 0
        adv_nativo: list[str] = []
        try:
            buffer.seek(0)
            texto_nativo, paginas, adv_nativo = _extraer_texto_pdf_nativo(buffer)
        except Exception as e:
            advertencias.append(f"PDF nativo falló ({e})")

        chars_nativo = len((texto_nativo or "").strip())
        if _texto_suficiente(texto_nativo, MIN_TEXTO_PDF_COMPLETO):
            return ResultadoIngesta(
                texto=texto_nativo,
                fuente="pdf_nativo",
                paginas=paginas,
                ocr_usado=False,
                calidad_imagen=1.0,
                advertencias=advertencias + adv_nativo,
            )

        advertencias.append(
            f"Texto nativo insuficiente ({chars_nativo} caracteres); aplicando OCR"
        )
        buffer.seek(0)
        try:
            texto_ocr, paginas_ocr, calidad, adv_ocr, fuente_ocr = _extraer_texto_ocr(
                buffer, nombre_archivo
            )
            partes = [t.strip() for t in (texto_nativo, texto_ocr) if t and t.strip()]
            texto_final = "\n\n".join(partes)
            return ResultadoIngesta(
                texto=texto_final,
                fuente=fuente_ocr if texto_ocr.strip() else "pdf_nativo_parcial",
                paginas=max(paginas, paginas_ocr),
                ocr_usado=bool(texto_ocr.strip()),
                calidad_imagen=calidad if texto_ocr.strip() else 0.85,
                advertencias=advertencias + adv_nativo + adv_ocr,
            )
        except Exception as e:
            if texto_nativo.strip():
                advertencias.append(f"OCR falló ({e}); usando solo texto nativo parcial")
                return ResultadoIngesta(
                    texto=texto_nativo,
                    fuente="pdf_nativo_parcial",
                    paginas=paginas,
                    ocr_usado=False,
                    calidad_imagen=0.6,
                    advertencias=advertencias + adv_nativo,
                )
            return ResultadoIngesta(
                texto="",
                fuente="vacio",
                paginas=0,
                ocr_usado=False,
                calidad_imagen=0.0,
                advertencias=advertencias,
                errores=[f"Error en OCR: {e}"],
            )

    # ── Imagen directa ──
    if nombre_lower.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
        fileobj.seek(0)
        try:
            texto, paginas, calidad, adv, fuente_ocr = _extraer_texto_ocr(fileobj, nombre_archivo)
            return ResultadoIngesta(
                texto=texto, fuente=fuente_ocr, paginas=paginas,
                ocr_usado=True, calidad_imagen=calidad,
                advertencias=advertencias + adv,
            )
        except Exception as e:
            return ResultadoIngesta(
                texto="", fuente="vacio", paginas=0, ocr_usado=False,
                calidad_imagen=0.0, errores=[f"Error en OCR de imagen: {e}"],
            )

    return ResultadoIngesta(
        texto="", fuente="vacio", paginas=0, ocr_usado=False,
        calidad_imagen=0.0, errores=[f"Formato de archivo no soportado: {nombre_archivo}"],
    )
